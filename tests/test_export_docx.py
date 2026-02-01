"""Unit tests for DOCXExporter."""


from docx import Document

from oeapp.services.export_docx import DOCXExporter
from tests.conftest import create_test_project, create_test_sentence



class TestDOCXExporter:
    """Test cases for DOCXExporter."""

    def test_export_creates_document(self, db_session, tmp_path):
        """Test export() creates a DOCX file."""
        project = create_test_project(db_session, name="Test Project", text="Se cyning.")

        exporter = DOCXExporter()
        output_path = tmp_path / "test.docx"

        result = exporter.export(project.id, output_path)

        assert result is True
        assert output_path.exists()

    def test_export_returns_false_when_project_not_found(self, db_session, tmp_path):
        """Test export() returns False when project doesn't exist."""
        exporter = DOCXExporter()
        output_path = tmp_path / "test.docx"

        result = exporter.export(99999, output_path)

        assert result is False

    def test_export_includes_project_title(self, db_session, tmp_path):
        """Test export() includes project name as heading."""
        project = create_test_project(
            db_session,
            name="My Test Project",
            text="",
            source="Test Source",
            translator="Test Translator",
            notes="Test Notes",
        )

        exporter = DOCXExporter()
        output_path = tmp_path / "test.docx"

        exporter.export(project.id, output_path)

        doc = Document(str(output_path))
        # First paragraph should be the heading
        assert len(doc.paragraphs) > 0
        assert doc.paragraphs[0].text == "My Test Project"

        # Check metadata
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Source: Test Source" in full_text
        assert "Translator: Test Translator" in full_text
        assert "Notes: Test Notes" in full_text

    def test_export_includes_sentence_numbers(self, db_session, tmp_path):
        """Test export() includes paragraph and sentence numbers."""
        project = create_test_project(db_session, name="Test", text="Se cyning. Þæt scip.")

        exporter = DOCXExporter()
        output_path = tmp_path / "test.docx"

        exporter.export(project.id, output_path)

        doc = Document(str(output_path))
        text = "\n".join([para.text for para in doc.paragraphs])

        # Should contain paragraph and sentence number markers
        assert "¶[" in text
        assert "S[" in text

    def test_export_includes_paragraph_breaks(self, db_session, tmp_path):
        """Test export() adds extra blank lines for paragraph starts."""
        project = create_test_project(db_session, name="Test", text="")
        
        from oeapp.models.chapter import Chapter
        from oeapp.models.section import Section
        from oeapp.models.paragraph import Paragraph
        
        chapter = Chapter(project_id=project.id, number=1)
        db_session.add(chapter)
        db_session.flush()
        section = Section(chapter_id=chapter.id, number=1)
        db_session.add(section)
        db_session.flush()
        
        p1 = Paragraph(section_id=section.id, order=1)
        p2 = Paragraph(section_id=section.id, order=2)
        db_session.add(p1)
        db_session.add(p2)
        db_session.flush()

        # Create sentences with paragraph breaks
        sentence1 = create_test_sentence(
            db_session, project_id=project.id, text="First paragraph.", display_order=1, paragraph_id=p1.id
        )
        sentence2 = create_test_sentence(
            db_session, project_id=project.id, text="Second paragraph.", display_order=2, paragraph_id=p2.id
        )

        exporter = DOCXExporter()
        output_path = tmp_path / "test.docx"

        exporter.export(project.id, output_path)

        doc = Document(str(output_path))
        # Should have multiple paragraphs (title + sentences + breaks)
        assert len(doc.paragraphs) > 2

    def test_export_includes_translation(self, db_session, tmp_path):
        """Test export() includes modern translation when available."""
        project = create_test_project(db_session, name="Test", text="Se cyning.")

        # Get the sentence and add translation
        sentence = project.sentences[0]
        sentence.text_modern = "The king"
        sentence.save()

        exporter = DOCXExporter()
        output_path = tmp_path / "test.docx"

        exporter.export(project.id, output_path)

        doc = Document(str(output_path))
        text = "\n".join([para.text for para in doc.paragraphs])

        assert "The king" in text

    def test_export_handles_missing_translation(self, db_session, tmp_path):
        """Test export() handles sentences without translation."""
        project = create_test_project(db_session, name="Test", text="Se cyning.")

        exporter = DOCXExporter()
        output_path = tmp_path / "test.docx"

        exporter.export(project.id, output_path)

        doc = Document(str(output_path))
        # Should still create document successfully
        assert len(doc.paragraphs) > 0

    def test_export_with_annotations_includes_superscripts(self, db_session, tmp_path):
        """Test export() includes superscript POS annotations."""
        project = create_test_project(db_session, name="Test", text="Se cyning")

        sentence = project.sentences[0]
        token = sentence.tokens[0]

        # Create annotation with POS
        if token.annotation:
            annotation = token.annotation
        else:
            from oeapp.models.annotation import Annotation
            annotation = Annotation(token_id=token.id)
            db_session.add(annotation)
            db_session.flush()

        annotation.pos = "R"
        annotation.pronoun_type = "d"
        annotation.save()

        exporter = DOCXExporter()
        output_path = tmp_path / "test.docx"

        exporter.export(project.id, output_path)

        doc = Document(str(output_path))
        # Find paragraph with Old English text
        oe_para = None
        for para in doc.paragraphs:
            if any("Se" in run.text for run in para.runs):
                oe_para = para
                break

        assert oe_para is not None
        # Check for superscript runs
        has_superscript = any(run.font.superscript for run in oe_para.runs)
        assert has_superscript

    def test_export_with_annotations_includes_subscripts(self, db_session, tmp_path):
        """Test export() includes subscript gender/context annotations."""
        project = create_test_project(db_session, name="Test", text="Se cyning")

        sentence = project.sentences[0]
        token = sentence.tokens[0]

        # Create annotation with gender
        if token.annotation:
            annotation = token.annotation
        else:
            from oeapp.models.annotation import Annotation
            annotation = Annotation(token_id=token.id)
            db_session.add(annotation)
            db_session.flush()

        annotation.pos = "N"
        annotation.gender = "m"
        annotation.case = "n"
        annotation.save()

        exporter = DOCXExporter()
        output_path = tmp_path / "test.docx"

        exporter.export(project.id, output_path)

        doc = Document(str(output_path))
        # Find paragraph with Old English text
        oe_para = None
        for para in doc.paragraphs:
            if any("cyning" in run.text for run in para.runs):
                oe_para = para
                break

        assert oe_para is not None
        # Check for subscript runs
        has_subscript = any(run.font.subscript for run in oe_para.runs)
        assert has_subscript

    def test_export_with_notes_includes_notes(self, db_session, tmp_path):
        """Test export() includes notes for sentences."""
        project = create_test_project(db_session, name="Test", text="Se cyning")

        sentence = project.sentences[0]
        token = sentence.tokens[0]

        # Create a note
        from oeapp.models.note import Note
        note = Note(
            sentence_id=sentence.id,
            start_token=token.id,
            end_token=token.id,
            note_text_md="This is a test note"
        )
        note.save()

        exporter = DOCXExporter()
        output_path = tmp_path / "test.docx"

        exporter.export(project.id, output_path)

        doc = Document(str(output_path))
        text = "\n".join([para.text for para in doc.paragraphs])

        assert "This is a test note" in text
        assert "1." in text  # Note numbering

    def test_export_with_multiple_notes_orders_correctly(self, db_session, tmp_path):
        """Test export() orders multiple notes by token position."""
        project = create_test_project(db_session, name="Test", text="Se cyning fēoll")

        sentence = project.sentences[0]
        tokens = list(sentence.tokens)

        # Create notes in reverse order
        from oeapp.models.note import Note
        note2 = Note(
            sentence_id=sentence.id,
            start_token=tokens[2].id,
            end_token=tokens[2].id,
            note_text_md="Note on third token"
        )
        note1 = Note(
            sentence_id=sentence.id,
            start_token=tokens[0].id,
            end_token=tokens[0].id,
            note_text_md="Note on first token"
        )
        note2.save(commit=False)
        note1.save()

        exporter = DOCXExporter()
        output_path = tmp_path / "test.docx"

        exporter.export(project.id, output_path)

        doc = Document(str(output_path))
        text = "\n".join([para.text for para in doc.paragraphs])

        # First note should appear before second note
        first_pos = text.find("Note on first token")
        second_pos = text.find("Note on third token")
        assert first_pos < second_pos

    def test_export_empty_project_creates_document(self, db_session, tmp_path):
        """Test export() creates document even for empty project."""
        project = create_test_project(db_session, name="Empty Project", text="")

        exporter = DOCXExporter()
        output_path = tmp_path / "test.docx"

        result = exporter.export(project.id, output_path)

        assert result is True
        assert output_path.exists()

        doc = Document(str(output_path))
        # Should have at least the title
        assert len(doc.paragraphs) > 0
        assert doc.paragraphs[0].text == "Empty Project"

    def test_setup_document_styles_sets_margins(self, db_session):
        """Test _setup_document_styles() sets document margins."""
        from docx import Document as DocxDocument

        exporter = DOCXExporter()
        doc = DocxDocument()

        exporter._setup_document_styles(doc)

        section = doc.sections[0]
        assert section.top_margin.inches == 1.0
        assert section.left_margin.inches == 1.0
        assert section.right_margin.inches == 1.0
        assert section.bottom_margin.inches == 1.0

    def test_export_handles_sentence_without_tokens(self, db_session, tmp_path):
        """Test export() handles sentence with no tokens gracefully."""
        project = create_test_project(db_session, name="Test", text="")

        # Create sentence manually without tokens
        from oeapp.models.sentence import Sentence
        from oeapp.models.chapter import Chapter
        from oeapp.models.section import Section
        from oeapp.models.paragraph import Paragraph
        
        chapter = Chapter(project_id=project.id, number=1)
        db_session.add(chapter)
        db_session.flush()
        section = Section(chapter_id=chapter.id, number=1)
        db_session.add(section)
        db_session.flush()
        paragraph = Paragraph(section_id=section.id, order=1)
        db_session.add(paragraph)
        db_session.flush()

        sentence = Sentence(
            project_id=project.id,
            display_order=1,
            text_oe="Test sentence",
            paragraph_id=paragraph.id
        )
        db_session.add(sentence)
        db_session.commit()

        exporter = DOCXExporter()
        output_path = tmp_path / "test.docx"

        result = exporter.export(project.id, output_path)

        assert result is True
        assert output_path.exists()

    def test_export_handles_file_write_error(self, db_session, tmp_path, monkeypatch):
        """Test export() handles file write errors gracefully."""
        project = create_test_project(db_session, name="Test", text="Se cyning.")

        exporter = DOCXExporter()
        output_path = tmp_path / "test.docx"

        # Mock Document.save to raise OSError
        original_save = None
        from docx.document import Document as DocumentClass

        def mock_save(self, path):
            raise OSError("Permission denied")

        # Patch the save method on the instance
        monkeypatch.setattr(DocumentClass, "save", mock_save)

        result = exporter.export(project.id, output_path)

        assert result is False

    def test_export_side_by_side_includes_metadata(self, db_session, tmp_path):
        """Test export_side_by_side() includes project metadata in header."""
        project = create_test_project(
            db_session,
            name="Side Project",
            text="Sentence one.",
            source="Side Source",
            translator="Side Translator",
            notes="Side Project Notes"
        )

        exporter = DOCXExporter()
        output_path = tmp_path / "side_by_side.docx"

        result = exporter.export_side_by_side(project.id, output_path)

        assert result is True
        assert output_path.exists()

        doc = Document(str(output_path))
        full_text = "\n".join([p.text for p in doc.paragraphs])

        assert "Translation: Side Project" in full_text
        assert "Source: Side Source" in full_text
        assert "Translator: Side Translator" in full_text
        assert "Project Notes: Side Project Notes" in full_text

    def test_export_side_by_side_includes_footer(self, db_session, tmp_path):
        """Test export_side_by_side() includes footer with metadata."""
        project = create_test_project(
            db_session,
            name="Footer Project",
            text="Sentence one.",
            translator="Footer Translator",
        )

        exporter = DOCXExporter()
        output_path = tmp_path / "footer_test.docx"

        exporter.export_side_by_side(project.id, output_path)

        doc = Document(str(output_path))
        section = doc.sections[0]
        footer = section.footer
        footer_text = "".join(p.text for p in footer.paragraphs)

        assert "Footer Project" in footer_text
        assert "Footer Translator" in footer_text

    def test_export_side_by_side_includes_horizontal_rule(self, db_session, tmp_path):
        """Test export_side_by_side() includes horizontal rule."""
        project = create_test_project(db_session, name="HR Project", text="One.")

        exporter = DOCXExporter()
        output_path = tmp_path / "hr_test.docx"

        exporter.export_side_by_side(project.id, output_path)

        doc = Document(str(output_path))
        # Find paragraph with border
        has_hr = False
        from oeapp.services.export_docx import qn

        for p in doc.paragraphs:
            p_pr = p._p.pPr
            if p_pr is not None:
                p_bdr = p_pr.find(qn("w:pBdr"))
                if p_bdr is not None and p_bdr.find(qn("w:bottom")) is not None:
                    has_hr = True
                    break
        assert has_hr

    def test_export_side_by_side_includes_vertical_rule(self, db_session, tmp_path):
        """Test export_side_by_side() includes vertical rule in table."""
        project = create_test_project(db_session, name="VR Project", text="One.")

        exporter = DOCXExporter()
        output_path = tmp_path / "vr_test.docx"

        exporter.export_side_by_side(project.id, output_path)

        doc = Document(str(output_path))
        table = doc.tables[0]
        # Check first cell's right border
        oe_cell = table.rows[0].cells[0]
        tc_pr = oe_cell._tc.tcPr
        from oeapp.services.export_docx import qn

        tc_borders = tc_pr.find(qn("w:tcBorders"))
        assert tc_borders is not None
        right_border = tc_borders.find(qn("w:right"))
        assert right_border is not None
        assert right_border.get(qn("w:color")) == "D3D3D3"

    def test_export_side_by_side_includes_footnotes_structure(
        self, db_session, tmp_path
    ):
        """Test export_side_by_side() creates manual hyperlinked endnotes."""
        project = create_test_project(
            db_session, name="Footnote Project", text="Se cyning."
        )
        sentence = project.sentences[0]
        token = sentence.tokens[0]

        from oeapp.models.note import Note

        note = Note(
            sentence_id=sentence.id,
            start_token=token.id,
            end_token=token.id,
            note_text_md="Manual note text",
        )
        note.save()

        exporter = DOCXExporter()
        output_path = tmp_path / "manual_footnote_test.docx"

        exporter.export_side_by_side(project.id, output_path)

        doc = Document(str(output_path))
        from oeapp.services.export_docx import qn

        # Check for hyperlinks in the main document
        table = doc.tables[0]
        oe_cell = table.rows[0].cells[0]
        xml_text = oe_cell._tc.xml
        assert "w:hyperlink" in xml_text
        assert 'w:anchor="note_1"' in xml_text

        # Check for bookmarks in the endnotes section
        full_xml = doc._element.xml
        assert "w:bookmarkStart" in full_xml
        assert 'w:name="note_1"' in full_xml
        assert "Manual note text" in full_xml

    def test_export_side_by_side_is_single_row(self, db_session, tmp_path):
        """Test export_side_by_side() uses a single-row table for flow."""
        project = create_test_project(
            db_session, name="Flow Project", text="First para. Second para."
        )
        # Create a second paragraph and move the second sentence to it
        from oeapp.models.paragraph import Paragraph
        p2 = Paragraph(section_id=project.chapters[0].sections[0].id, order=2)
        db_session.add(p2)
        db_session.flush()
        
        project.sentences[1].paragraph_id = p2.id
        db_session.commit()

        exporter = DOCXExporter()
        output_path = tmp_path / "single_row_test.docx"

        exporter.export_side_by_side(project.id, output_path)

        doc = Document(str(output_path))
        table = doc.tables[0]
        # Should have exactly one row despite having two paragraphs
        assert len(table.rows) == 1

        # Should have multiple paragraphs in the cell
        oe_cell = table.rows[0].cells[0]
        # At least 2 paragraphs (the two sentences)
        assert len(oe_cell.paragraphs) >= 2
