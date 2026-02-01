"""DOCX export service for Ænglisc Toolkit."""

from typing import TYPE_CHECKING, cast

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from oeapp.mixins import AnnotationTextualMixin, TokenOccurrenceMixin
from oeapp.models.mixins import SessionMixin
from oeapp.models.project import Project

if TYPE_CHECKING:
    from pathlib import Path

    from docx.document import Document as DocumentObject
    from docx.text.run import Run

    from oeapp.models.annotation import Annotation
    from oeapp.models.sentence import Sentence
    from oeapp.models.token import Token


class DOCXExporter(SessionMixin, AnnotationTextualMixin, TokenOccurrenceMixin):
    """
    Exports annotated Old English text to DOCX format.

    """

    def __init__(self) -> None:
        """
        Initialize exporter.

        """
        self.session = self._get_session()

    def _add_horizontal_rule(self, doc: "DocumentObject") -> None:
        """
        Add a horizontal rule to the document.

        Args:
            doc: The document to add the rule to

        """
        p = doc.add_paragraph()
        p_pr = p._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    def _set_cell_border(self, cell, **kwargs) -> None:
        """
        Set cell borders.

        Args:
            cell: The cell to modify
            **kwargs: Border settings (top, bottom, left, right, start, end)

        """
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()

        tc_borders = tc_pr.find(qn("w:tcBorders"))
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)

        for edge in ("top", "start", "bottom", "end", "left", "right"):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = f"w:{edge}"
                element = tc_borders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tc_borders.append(element)

                for key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))

    def _add_note_reference(self, paragraph, note_num: int) -> None:
        """
        Add a manual hyperlinked note reference to a paragraph.

        Args:
            paragraph: The paragraph to add the reference to
            note_num: The number of the note

        """
        # Create hyperlink element
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), f"note_{note_num}")
        hyperlink.set(qn("w:history"), "1")

        # Create run for the number
        new_run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        r_style = OxmlElement("w:rStyle")
        r_style.set(qn("w:val"), "FootnoteReference")
        r_pr.append(r_style)

        # Also make it superscript explicitly
        vert_align = OxmlElement("w:vertAlign")
        vert_align.set(qn("w:val"), "superscript")
        r_pr.append(vert_align)

        new_run.append(r_pr)

        text_el = OxmlElement("w:t")
        text_el.text = str(note_num)
        new_run.append(text_el)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def _add_note_anchor(
        self, doc: "DocumentObject", note_num: int, note_text: str, bookmark_id: int
    ) -> None:
        """
        Add a manual endnote entry with a bookmark anchor.

        Args:
            doc: The document to add the entry to
            note_num: The number of the note
            note_text: The text of the note
            bookmark_id: Unique ID for the bookmark

        """
        p = doc.add_paragraph()

        # Bookmark start
        bm_start = OxmlElement("w:bookmarkStart")
        bm_start.set(qn("w:id"), str(bookmark_id))
        bm_start.set(qn("w:name"), f"note_{note_num}")
        p._p.append(bm_start)

        # Note content
        run = p.add_run(f"{note_num}. ")
        run.bold = True
        p.add_run(note_text)

        # Bookmark end
        bm_end = OxmlElement("w:bookmarkEnd")
        bm_end.set(qn("w:id"), str(bookmark_id))
        p._p.append(bm_end)

    def _add_page_number(self, run: "Run") -> None:
        """
        Add a page number field to a run.

        Args:
            run: The run to add the page number to

        """
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = "PAGE"

        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")

        run._r.append(fld_char1)
        run._r.append(instr_text)
        run._r.append(fld_char2)

    def export(self, project_id: int, output_path: "Path") -> bool:
        """
        Export project to DOCX file.

        Args:
            project_id: Project ID to export
            output_path: Path to output DOCX file

        Returns:
            True if successful, False otherwise

        """
        doc: DocumentObject = Document()
        self._setup_document_styles(doc)
        project = Project.get(project_id)
        if project is None:
            return False

        # Add title
        doc.add_heading(project.name, level=1)

        # Add metadata
        if project.source:
            source_para = doc.add_paragraph()
            source_para.add_run("Source: ").bold = True
            source_para.add_run(project.source)

        if project.translator:
            translator_para = doc.add_paragraph()
            translator_para.add_run("Translator: ").bold = True
            translator_para.add_run(project.translator)

        if project.notes:
            notes_para = doc.add_paragraph()
            notes_para.add_run("Notes: ").bold = True
            notes_para.add_run(project.notes)

        doc.add_paragraph()  # Blank line after metadata

        for sentence in project.sentences:
            text_modern = sentence.text_modern

            # Add paragraph break if this sentence starts a paragraph
            is_paragraph_start = False
            if sentence.paragraph:
                p_sentences = sorted(sentence.paragraph.sentences, key=lambda s: s.display_order)
                if p_sentences and p_sentences[0].id == sentence.id:
                    is_paragraph_start = True

            if is_paragraph_start:
                doc.add_paragraph()  # Extra blank line for paragraph break

            # Add sentence number with paragraph and sentence numbers
            sentence_num_para = doc.add_paragraph()
            paragraph_num = sentence.paragraph.order if sentence.paragraph else 0
            # Calculate sentence number in paragraph
            sentence_num = 1
            if sentence.paragraph:
                p_sentences = sorted(sentence.paragraph.sentences, key=lambda s: s.display_order)
                for i, s in enumerate(p_sentences, 1):
                    if s.id == sentence.id:
                        sentence_num = i
                        break

            sentence_num_run = sentence_num_para.add_run(
                f"¶[{paragraph_num}] S[{sentence_num}] "
            )
            sentence_num_run.bold = True

            # Build Old English sentence with annotations
            self._add_oe_sentence_with_annotations(doc, sentence)

            # Add translation
            if text_modern:
                translation_para = doc.add_paragraph()
                translation_run = translation_para.add_run(text_modern)
                translation_run.font.size = Pt(12)
                translation_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            else:
                # Empty translation paragraph
                doc.add_paragraph()

            # Blank line
            doc.add_paragraph()

            # Add notes
            self._add_notes(doc, sentence)

            # Blank line between sentences
            doc.add_paragraph()

        try:
            doc.save(str(output_path))
        except OSError as e:
            print(f"Export error: {e}")  # noqa: T201
            return False
        else:
            return True

    def export_side_by_side(self, project_id: int, output_path: "Path") -> bool:  # noqa: PLR0912, PLR0915
        """
        Export project to DOCX file in side-by-side landscape format.

        Args:
            project_id: Project ID to export
            output_path: Path to output DOCX file

        Returns:
            True if successful, False otherwise

        """
        doc: DocumentObject = Document()

        # Set landscape orientation
        section = doc.sections[0]
        # In python-docx, to swap to landscape we must swap height and width
        original_width = section.page_width
        original_height = section.page_height
        section.page_width = original_height
        section.page_height = original_width
        # landscape orientation is 1
        section.orientation = 1  # type: ignore[assignment]

        # Standard margins
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        project = Project.get(project_id)
        if project is None:
            return False

        doc.add_heading(f"Translation: {project.name}", level=1)
        if project.source:
            doc.add_paragraph(f"Source: {project.source}")
        if project.translator:
            doc.add_paragraph(f"Translator: {project.translator}")
        if project.notes:
            doc.add_paragraph(f"Project Notes: {project.notes}")

        # Add horizontal rule
        self._add_horizontal_rule(doc)

        # Add footer
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.add_run(f"{project.name} | ")
        if project.translator:
            footer_p.add_run(f"Translator: {project.translator} | ")
        self._add_page_number(footer_p.add_run())

        # Collect and number all notes project-wide
        project_notes = []
        note_num = 1
        note_map: dict[
            tuple[int, int], list[int]
        ] = {}  # (sentence_id, token_id) -> list of note numbers
        for sentence in project.sentences:
            for note in sentence.sorted_notes:
                project_notes.append((note_num, note))
                if note.end_token:
                    key = (sentence.id, note.end_token)
                    if key not in note_map:
                        note_map[key] = []
                    note_map[key].append(note_num)
                note_num += 1

        # Use a table for side-by-side alignment - SINGLE ROW for independent flow
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False

        # Page width is 11 inches, margins 0.5 each -> 10 inches available
        col_width = Inches(5.0)

        oe_cell = table.rows[0].cells[0]
        mode_cell = table.rows[0].cells[1]
        oe_cell.width = col_width
        mode_cell.width = col_width

        # Add vertical rule between columns (right border of OE cell)
        self._set_cell_border(
            oe_cell,
            right={"sz": 4, "val": "single", "color": "D3D3D3", "space": "0"},
        )

        current_oe_p = oe_cell.paragraphs[0]
        current_mode_p = mode_cell.paragraphs[0]

        for sentence in project.sentences:
            # Check if this sentence is the first in its paragraph
            is_paragraph_start = False
            if sentence.paragraph:
                p_sentences = sorted(sentence.paragraph.sentences, key=lambda s: s.display_order)
                if p_sentences and p_sentences[0].id == sentence.id:
                    is_paragraph_start = True

            if is_paragraph_start:
                # If not the first sentence, create new paragraphs in both cells
                if sentence != project.sentences[0]:
                    current_oe_p = oe_cell.add_paragraph()
                    current_mode_p = mode_cell.add_paragraph()
            else:
                # Add a space between sentences in the same paragraph
                # But only if it's not the very first sentence
                if sentence != project.sentences[0]:
                    current_oe_p.add_run(" ")
                    current_mode_p.add_run(" ")

            # Add OE text with manual hyperlinked note references
            tokens, token_id_to_start = sentence.sorted_tokens
            text = sentence.text_oe
            last_pos = 0
            for token in tokens:
                token_start = token_id_to_start[token.id]
                if token_start > last_pos:
                    current_oe_p.add_run(text[last_pos:token_start])

                current_oe_p.add_run(token.surface)

                # Add manual note references
                key = (sentence.id, token.id)
                if key in note_map:
                    for n_num in note_map[key]:
                        self._add_note_reference(current_oe_p, n_num)

                last_pos = token_start + len(token.surface)

            if last_pos < len(text):
                current_oe_p.add_run(text[last_pos:])

            # Add ModE text
            mode_run = current_mode_p.add_run(sentence.text_modern or "[...]")
            if not sentence.text_modern:
                mode_run.font.italic = True

        # Add notes at the end of the document (Endnotes Section)
        if project_notes:
            doc.add_page_break()
            doc.add_heading("Notes", level=2)

            # Bookmark ID counter (start at a reasonable number)
            bookmark_id = 1000
            for n_num, note in project_notes:
                # Get token text for the note
                token_id_to_order = {t.id: t.order_index for t in note.sentence.tokens}
                token_text = self._get_note_token_text(
                    note, note.sentence, token_id_to_order
                )

                if token_text:
                    note_text = f'"{token_text}" - {note.note_text_md}'
                else:
                    note_text = note.note_text_md

                self._add_note_anchor(doc, n_num, note_text, bookmark_id)
                bookmark_id += 1

        try:
            doc.save(str(output_path))
        except OSError as e:
            print(f"Export error: {e}")  # noqa: T201
            return False
        else:
            return True

    def _setup_document_styles(self, doc: "DocumentObject") -> None:
        """
        Set up document styles.

        Args:
            doc: Document to set up

        """
        # Styles are already available in python-docx
        # Title, Body, Default are standard styles
        # Set the top, left, right, bottom margins to 1 inch
        doc.sections[0].top_margin = Inches(1)
        doc.sections[0].left_margin = Inches(1)
        doc.sections[0].right_margin = Inches(1)
        doc.sections[0].bottom_margin = Inches(1)

    def _set_run_position(self, run, position_half_points: int) -> None:
        """
        Set the vertical position of a run via XML manipulation.

        This adjusts the baseline position of text (e.g., superscripts).
        Position is specified in half-points (1/144 of an inch).

        Args:
            run: The run to modify
            position_half_points: Position offset in half-points
                (positive values raise the text, negative lower it)

        """
        # Access the run's XML element
        r = run._element

        # Find or create the <w:rPr> element (run properties)
        r_pr = r.find(qn("w:rPr"))
        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            r.insert(0, r_pr)

        # Find or create the <w:position> element
        position = r_pr.find(qn("w:position"))
        if position is None:
            position = OxmlElement("w:position")
            r_pr.append(position)

        # Set the position value (in half-points)
        position.set(qn("w:val"), str(position_half_points))

    def _add_oe_sentence_with_annotations(
        self,
        doc: "DocumentObject",
        sentence: "Sentence",
    ) -> None:
        """
        Add Old English sentence with superscript/subscript annotations.

        Uses the original sentence.text_oe to preserve all punctuation and spacing,
        then adds annotations (pos, gender, context) for each token.

        Args:
            doc: Document to add to
            sentence: Sentence to add

        """
        # Build paragraph with annotations
        para = doc.add_paragraph()  # type: ignore[attr-defined]

        # Use original sentence text to preserve punctuation
        text = sentence.text_oe
        tokens = list(sentence.tokens)

        if not tokens:
            # No tokens, just add the text as-is
            para.add_run(text)
            return

        # Sort tokens by order_index to process them in order
        sorted_tokens = sorted(tokens, key=lambda t: t.order_index)

        # Find token positions in the original text
        # (start, end, token) positions
        token_positions: list[tuple[int, int, Token]] = []
        used_positions: set[tuple[int, int]] = set()

        for token in sorted_tokens:
            if not token.surface:
                continue
            token_start = self._find_token_occurrence(text, token, tokens)
            if token_start is not None:
                token_end = token_start + len(token.surface)
                position_key = (token_start, token_end)

                # Only add if this position hasn't been used yet
                if position_key not in used_positions:
                    token_positions.append((token_start, token_end, token))
                    used_positions.add(position_key)

        # Sort by position
        token_positions.sort(key=lambda x: x[0])

        # Build document by preserving text between tokens
        last_pos = 0
        for token_start, token_end, token in token_positions:
            # Skip if this token overlaps with a previous one
            if token_start < last_pos:
                continue

            # Add text before token (preserving punctuation and spacing)
            if token_start > last_pos:
                para.add_run(text[last_pos:token_start])

            # Add token with annotations
            annotation = token.annotation

            # POS label (superscript)
            pos_label = self.format_pos(cast("Annotation", annotation))
            if pos_label:
                pos_run = para.add_run(pos_label)
                pos_run.font.size = Pt(8)
                pos_run.font.superscript = True
                # Raise baseline to height of capital letters
                # (6 points = 12 half-points)
                self._set_run_position(pos_run, 4)
                pos_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

            # Gender label (subscript)
            gender_label = self.format_gender(cast("Annotation", annotation))
            if gender_label:
                gender_run = para.add_run(gender_label)
                gender_run.font.size = Pt(8)
                gender_run.font.subscript = True
                gender_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

            # Add word
            word_run = para.add_run(token.surface)
            word_run.font.size = Pt(12)

            # Context label (subscript)
            context_label = self.format_context(cast("Annotation", annotation))
            if context_label:
                context_run = para.add_run(context_label)
                context_run.font.size = Pt(8)
                context_run.font.subscript = True
                context_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

            last_pos = token_end

        # Add remaining text after last token
        if last_pos < len(text):
            para.add_run(text[last_pos:])

    def _add_notes(self, doc: "DocumentObject", sentence: "Sentence") -> None:
        """
        Add notes for a sentence.

        Notes are sorted by their position in the sentence (by start token
        order_index) and numbered accordingly.

        Args:
            doc: Document to add to
            sentence: Sentence to add notes for

        """
        if not sentence.notes:
            return

        # Sort notes by token position in sentence (earlier tokens = lower numbers)
        notes = self._sort_notes_by_position(sentence)

        # Build token ID to order_index mapping for token lookups
        token_id_to_order: dict[int, int] = {}
        for token in sentence.tokens:
            if token.id:
                token_id_to_order[token.id] = token.order_index

        # Display each note with dynamic numbering (1-based index)
        for note_idx, note in enumerate(notes, start=1):
            # Get token text for the note
            token_text = self._get_note_token_text(note, sentence, token_id_to_order)

            # Format note: "1. "quoted tokens" in italics - note text"
            if token_text:
                note_line = f'{note_idx}. "{token_text}" - {note.note_text_md}'
            else:
                note_line = f"{note_idx}. {note.note_text_md}"

            note_para = doc.add_paragraph()
            note_run = note_para.add_run(note_line)
            note_run.font.size = Pt(10)
            note_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    def _sort_notes_by_position(self, sentence: "Sentence") -> list:
        """
        Sort notes by their position in the sentence (by start token order_index).

        Args:
            sentence: Sentence to get notes from

        Returns:
            Sorted list of notes

        """
        # Build token ID to order_index mapping
        token_id_to_order: dict[int, int] = {}
        for token in sentence.tokens:
            if token.id:
                token_id_to_order[token.id] = token.order_index

        def get_note_position(note) -> int:
            """Get position of note in sentence based on start token."""
            if note.start_token and note.start_token in token_id_to_order:
                return token_id_to_order[note.start_token]
            # Fallback to end_token if start_token not found
            if note.end_token and note.end_token in token_id_to_order:
                return token_id_to_order[note.end_token]
            # Fallback to very high number if neither found
            return 999999

        # Sort notes by position
        return sorted(sentence.notes, key=get_note_position)

    def _get_note_token_text(
        self,
        note,
        sentence: "Sentence",
        token_id_to_order: dict[int, int],  # noqa: ARG002
    ) -> str:
        """
        Get token text for a note.

        Args:
            note: Note to get tokens for
            sentence: Sentence containing the tokens
            token_id_to_order: Map of token ID to order_index

        Returns:
            Token text string (space-separated tokens)

        """
        if not note.start_token or not note.end_token:
            return ""

        # Find tokens by ID
        start_token = None
        end_token = None
        for token in sentence.tokens:
            if token.id == note.start_token:
                start_token = token
            if token.id == note.end_token:
                end_token = token

        if not start_token or not end_token:
            return ""

        # Get all tokens in range
        tokens_in_range = []
        in_range = False
        for token in sorted(sentence.tokens, key=lambda t: t.order_index):
            if token.id == start_token.id:
                in_range = True
            if in_range:
                tokens_in_range.append(token.surface)
            if token.id == end_token.id:
                break

        return " ".join(tokens_in_range)
